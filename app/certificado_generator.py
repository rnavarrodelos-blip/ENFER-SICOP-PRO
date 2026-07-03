from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from app.db_init import get_session
from app.models import Certificado, Participante, Curso, Voucher
import os


class CertificadoGenerator:
    """Genera certificados en formato Word"""

    def __init__(self):
        self.session = get_session()
        if not os.path.exists('certificados'):
            os.makedirs('certificados')

    def generar_certificado_word(self, participante_id, curso_id):
        """Genera un certificado en Word para un participante"""
        participante = self.session.query(Participante).get(participante_id)
        curso = self.session.query(Curso).get(curso_id)
        voucher = self.session.query(Voucher).filter_by(
            participante_id=participante_id
        ).first()

        if not participante or not curso:
            return {"success": False, "mensaje": "Participante o curso no encontrado"}

        if not voucher or not voucher.verified:
            return {"success": False, "mensaje": "Voucher no verificado"}

        # Verificar si ya existe certificado
        certificado_existente = self.session.query(Certificado).filter_by(
            participante_id=participante_id,
            curso_id=curso_id
        ).first()

        if certificado_existente:
            return {"success": False, "mensaje": "Certificado ya existe para este participante", "duplicado": True}

        # Crear documento
        doc = Document()

        # Título
        titulo = doc.add_paragraph()
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        titulo_run = titulo.add_run("CERTIFICADO DE PARTICIPACIÓN")
        titulo_run.font.size = Pt(28)
        titulo_run.font.bold = True
        titulo_run.font.color.rgb = RGBColor(11, 94, 215)  # Color #0B5ED7

        # Subtítulo
        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitulo_run = subtitulo.add_run("ENFER-SICOP PRO")
        subtitulo_run.font.size = Pt(14)
        subtitulo_run.font.bold = True

        doc.add_paragraph()  # Espacio

        # Cuerpo
        cuerpo = doc.add_paragraph()
        cuerpo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cuerpo_text = f"""Se certifica que {participante.nombre} ha participado exitosamente en el curso:"""
        cuerpo_run = cuerpo.add_run(cuerpo_text)
        cuerpo_run.font.size = Pt(12)

        # Nombre del curso
        curso_para = doc.add_paragraph()
        curso_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        curso_run = curso_para.add_run(f'"{curso.nombre}"')
        curso_run.font.size = Pt(14)
        curso_run.font.bold = True
        curso_run.font.color.rgb = RGBColor(11, 94, 215)

        doc.add_paragraph()  # Espacio

        # Detalles
        detalles = doc.add_paragraph()
        detalles.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        detalles_text = f"""Duración: {curso.horas} horas
Fecha de emisión: {datetime.now().strftime('%d de %B de %Y')}"""
        detalles_run = detalles.add_run(detalles_text)
        detalles_run.font.size = Pt(11)

        doc.add_paragraph()  # Espacio
        doc.add_paragraph()  # Espacio

        # Número de certificado
        numero_cert = f"CERT-{datetime.now().strftime('%Y%m%d%H%M%S')}"
        numero_para = doc.add_paragraph()
        numero_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        numero_run = numero_para.add_run(f"Certificado N°: {numero_cert}")
        numero_run.font.size = Pt(10)
        numero_run.font.italic = True

        # Guardar documento
        archivo_nombre = f"CERTIFICADO_{participante.nombre.replace(' ', '_')}_{curso.nombre.replace(' ', '_')}.docx"
        archivo_path = os.path.join('certificados', archivo_nombre)
        doc.save(archivo_path)

        # Guardar en BD
        nuevo_certificado = Certificado(
            numero_certificado=numero_cert,
            participante_id=participante_id,
            curso_id=curso_id,
            archivo_path=archivo_path,
            estado="generado"
        )
        self.session.add(nuevo_certificado)
        self.session.commit()

        return {
            "success": True,
            "mensaje": "Certificado generado exitosamente",
            "archivo_path": archivo_path,
            "numero_certificado": numero_cert
        }

    def generar_certificados_lote(self, curso_id):
        """Genera certificados para todos los participantes de un curso con voucher verificado"""
        participantes = self.session.query(Participante).filter_by(curso_id=curso_id).all()
        resultados = []

        for participante in participantes:
            # Verificar que tenga voucher verificado
            voucher = self.session.query(Voucher).filter_by(
                participante_id=participante.id
            ).filter_by(verified=True).first()

            if voucher:
                resultado = self.generar_certificado_word(participante.id, curso_id)
                resultados.append({
                    "participante": participante.nombre,
                    "resultado": resultado
                })

        return resultados
